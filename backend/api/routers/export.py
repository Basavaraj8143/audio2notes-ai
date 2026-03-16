import io
import textwrap

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from docx import Document

from models.session import sessions

router = APIRouter()


@router.get("/{session_id}/pdf")
async def export_pdf(session_id: str):
    """Export notes as a formatted PDF."""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found.")
    data = sessions[session_id]
    buf = io.BytesIO()
    _build_pdf(buf, data)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="notes_{session_id[:8]}.pdf"'},
    )


@router.get("/{session_id}/docx")
async def export_docx(session_id: str):
    """Export notes as a formatted DOCX."""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found.")
    data = sessions[session_id]
    buf = io.BytesIO()
    _build_docx(buf, data)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="notes_{session_id[:8]}.docx"'},
    )


@router.get("/{session_id}/txt")
async def export_txt(session_id: str):
    """Export notes as plain text."""
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found.")
    text = sessions[session_id]["merged_notes"]
    return StreamingResponse(
        io.BytesIO(text.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="notes_{session_id[:8]}.txt"'},
    )


def _build_pdf(buf, data):
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=20, textColor=colors.HexColor("#6C63FF"))
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, textColor=colors.HexColor("#333"))
    body_style = styles["BodyText"]

    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    story = [Paragraph(f"Lecture Notes: {data['filename']}", title_style), Spacer(1, 0.5 * cm)]

    for i, chunk in enumerate(data["notes_chunks"]):
        story.append(Paragraph(f"Section {i + 1}: {chunk.get('topic', 'Unknown')}", h2_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
        story.append(Spacer(1, 0.2 * cm))

        for kp in chunk.get("key_points", []):
            story.append(Paragraph(f"• {kp}", body_style))

        if chunk.get("definitions"):
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("<b>Definitions:</b>", body_style))
            for term, defn in chunk["definitions"].items():
                story.append(Paragraph(f"<b>{term}</b>: {defn}", body_style))

        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(f"<i>{chunk.get('summary', '')}</i>", body_style))
        story.append(Spacer(1, 0.5 * cm))

    doc.build(story)


def _build_docx(buf, data):
    doc = Document()
    doc.add_heading(f"Lecture Notes: {data['filename']}", 0)
    for i, chunk in enumerate(data["notes_chunks"]):
        doc.add_heading(f"Section {i + 1}: {chunk.get('topic', '')}", level=1)
        for kp in chunk.get("key_points", []):
            doc.add_paragraph(kp, style="List Bullet")
        if chunk.get("definitions"):
            doc.add_heading("Definitions", level=2)
            for term, defn in chunk["definitions"].items():
                p = doc.add_paragraph()
                p.add_run(f"{term}: ").bold = True
                p.add_run(defn)
        doc.add_paragraph(chunk.get("summary", ""), style="Body Text")
        doc.add_paragraph()
    doc.save(buf)
